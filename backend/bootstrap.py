import os
import textwrap

BASE_DIR = r"C:\Users\KenJes\.gemini\antigravity\scratch\raygpt_mobile\backend"

FILES = {
    "requirements.txt": """
fastapi>=0.115.0
uvicorn[standard]>=0.32.0
pydantic>=2.10.0
pydantic-settings>=2.6.0
litellm>=1.50.0
google-genai>=1.0.0
langchain>=0.3.0
langchain-google-genai>=2.0.0
langgraph>=0.2.0
qdrant-client>=1.12.0
sentence-transformers>=3.3.0
mcp>=1.2.0
sqlalchemy>=2.0.0
asyncpg>=0.30.0
alembic>=1.14.0
redis>=5.2.0
google-api-python-client>=2.150.0
google-auth-oauthlib>=1.2.0
spotipy>=2.24.0
python-jose[cryptography]>=3.3.0
python-multipart>=0.0.12
websockets>=13.0
httpx>=0.28.0
python-dotenv>=1.0.0
aiofiles>=24.0.0
loguru>=0.7.0
PyPDF2>=3.0.0
python-docx>=1.1.0
    """,
    "Dockerfile": """
FROM python:3.12-slim

WORKDIR /app

RUN apt-get update && apt-get install -y --no-install-recommends \
    build-essential \
    && rm -rf /var/lib/apt/lists/*

COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

COPY . .

EXPOSE 8000

CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]
    """,
    "docker-compose.yml": """
version: '3.8'

services:
  backend:
    build: .
    ports:
      - "8000:8000"
    volumes:
      - .:/app
    env_file:
      - .env
    depends_on:
      - postgres
      - redis
      - qdrant
    networks:
      - raygpt_network

  postgres:
    image: postgres:16
    environment:
      POSTGRES_USER: raymundo
      POSTGRES_PASSWORD: password
      POSTGRES_DB: raygpt
    ports:
      - "5432:5432"
    volumes:
      - postgres_data:/var/lib/postgresql/data
    networks:
      - raygpt_network

  redis:
    image: redis:7
    ports:
      - "6379:6379"
    volumes:
      - redis_data:/data
    networks:
      - raygpt_network

  qdrant:
    image: qdrant/qdrant:latest
    ports:
      - "6333:6333"
    volumes:
      - qdrant_data:/qdrant/storage
    networks:
      - raygpt_network

volumes:
  postgres_data:
  redis_data:
  qdrant_data:

networks:
  raygpt_network:
    driver: bridge
    """,
    ".env.example": """
GEMINI_API_KEY=your_gemini_api_key_here
OPENAI_API_KEY=your_openai_api_key_here

DATABASE_URL=postgresql+asyncpg://raymundo:password@localhost:5432/raygpt
REDIS_URL=redis://localhost:6379/0
QDRANT_URL=http://localhost:6333

JWT_SECRET=your_super_secret_jwt_key_here

GOOGLE_CLIENT_ID=your_google_client_id
GOOGLE_CLIENT_SECRET=your_google_client_secret

SPOTIFY_CLIENT_ID=your_spotify_client_id
SPOTIFY_CLIENT_SECRET=your_spotify_client_secret
    """,
    "app/__init__.py": "",
    "app/main.py": """
import contextlib
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from app.api.v1 import chat, documents, conversations
from app.db.database import init_db
from app.db.redis_client import init_redis, close_redis
from app.core.mcp.client_manager import mcp_manager
from loguru import logger

@contextlib.asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Initializing RayGPT 2.0 Backend...")
    await init_db()
    await init_redis()
    await mcp_manager.initialize()
    logger.info("RayGPT 2.0 Backend started successfully.")
    yield
    logger.info("Shutting down RayGPT 2.0 Backend...")
    await close_redis()
    await mcp_manager.shutdown()
    logger.info("Shutdown complete.")

app = FastAPI(title="RayGPT 2.0 API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(chat.router, prefix="/api/v1/chat", tags=["Chat"])
app.include_router(documents.router, prefix="/api/v1/documents", tags=["Documents"])
app.include_router(conversations.router, prefix="/api/v1/conversations", tags=["Conversations"])

@app.get("/health")
async def health_check():
    return {"status": "ok", "version": "2.0.0"}
    """,
    "app/config.py": """
from pydantic_settings import BaseSettings, SettingsConfigDict
from pydantic import Field

class Settings(BaseSettings):
    gemini_api_key: str = Field(default="")
    openai_api_key: str = Field(default="")
    
    database_url: str = Field(default="postgresql+asyncpg://raymundo:password@localhost:5432/raygpt")
    redis_url: str = Field(default="redis://localhost:6379/0")
    qdrant_url: str = Field(default="http://localhost:6333")
    
    jwt_secret: str = Field(default="dev_secret")
    
    default_model: str = Field(default="gemini/gemini-2.5-flash")
    vision_model: str = Field(default="gemini/gemini-2.5-pro")
    embedding_model: str = Field(default="all-MiniLM-L6-v2")

    model_config = SettingsConfigDict(env_file=".env", env_file_encoding="utf-8", extra="ignore")

settings = Settings()
    """,
    "app/api/__init__.py": "",
    "app/api/v1/__init__.py": "",
    "app/api/v1/chat.py": """
from fastapi import APIRouter, WebSocket, WebSocketDisconnect, Depends
from typing import Dict
import json
from loguru import logger
from app.models.schemas import MessageCreate, MessageResponse, WebSocketMessage, WebSocketResponse
from app.services.memory import memory_service
from app.core.llm.router import llm_router
from app.core.rag.pipeline import rag_pipeline
from app.core.mcp.client_manager import mcp_manager
from app.core.llm.prompts import RAYMUNDO_SYSTEM_PROMPT_AMIGABLE, RAYMUNDO_SYSTEM_PROMPT_DIRECTO

router = APIRouter()

class ConnectionManager:
    def __init__(self):
        self.active_connections: list[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

manager = ConnectionManager()

@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            message_data = json.loads(data)
            msg = WebSocketMessage(**message_data)
            
            if msg.type == 'message':
                history = await memory_service.get_history(msg.conversation_id)
                
                # Setup context
                system_prompt = RAYMUNDO_SYSTEM_PROMPT_AMIGABLE if msg.mode == 'amigable' else RAYMUNDO_SYSTEM_PROMPT_DIRECTO
                messages = [{"role": "system", "content": system_prompt}]
                
                # Check RAG
                rag_results = await rag_pipeline.query(msg.content, msg.user_id)
                if rag_results:
                    context = "\\n".join([r.content for r in rag_results])
                    messages.append({"role": "system", "content": f"Contexto:\\n{context}"})
                
                for h in history:
                    messages.append({"role": h.role, "content": h.content})
                
                messages.append({"role": "user", "content": msg.content})
                
                full_response = ""
                async for chunk in llm_router.stream_completion(messages):
                    full_response += chunk
                    await websocket.send_json(WebSocketResponse(type="token", content=chunk).model_dump())
                
                # Save messages
                await memory_service.save_message(msg.conversation_id, "user", msg.content)
                await memory_service.save_message(msg.conversation_id, "assistant", full_response)
                
                await websocket.send_json(WebSocketResponse(type="message_complete", content=full_response, message_id="temp-id").model_dump())
                
    except WebSocketDisconnect:
        manager.disconnect(websocket)
    except Exception as e:
        logger.error(f"WebSocket error: {e}")
        await websocket.send_json(WebSocketResponse(type="error", content=str(e)).model_dump())

@router.post("/message", response_model=MessageResponse)
async def send_message_rest(message: MessageCreate):
    # Fallback REST implementation
    history = await memory_service.get_history(message.conversation_id)
    messages = [{"role": h.role, "content": h.content} for h in history]
    messages.append({"role": "user", "content": message.content})
    
    response_text = await llm_router.completion(messages)
    
    await memory_service.save_message(message.conversation_id, "user", message.content)
    saved_msg = await memory_service.save_message(message.conversation_id, "assistant", response_text)
    
    return MessageResponse(
        id=saved_msg.id,
        content=saved_msg.content,
        role=saved_msg.role,
        timestamp=saved_msg.timestamp,
        tool_calls=[]
    )
    """,
    "app/api/v1/documents.py": """
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import List
from app.models.schemas import DocumentResponse, DocumentUploadResponse
from app.core.rag.pipeline import rag_pipeline
import os
import aiofiles

router = APIRouter()

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    user_id: str = Form(...)
):
    temp_path = f"/tmp/{file.filename}"
    async with aiofiles.open(temp_path, 'wb') as out_file:
        content = await file.read()
        await out_file.write(content)
        
    try:
        metadata = {"filename": file.filename, "user_id": user_id}
        doc_id = await rag_pipeline.ingest_document(temp_path, user_id, metadata)
        return DocumentUploadResponse(
            id=doc_id,
            filename=file.filename,
            file_type=file.filename.split('.')[-1],
            chunk_count=0, # Simplified for now
            created_at="now"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@router.get("", response_model=List[DocumentResponse])
async def list_documents():
    return []

@router.delete("/{doc_id}")
async def delete_document(doc_id: str):
    success = await rag_pipeline.delete_document(doc_id)
    if not success:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}

@router.get("/{doc_id}", response_model=DocumentResponse)
async def get_document(doc_id: str):
    raise HTTPException(status_code=501, detail="Not implemented")
    """,
    "app/api/v1/conversations.py": """
from fastapi import APIRouter, HTTPException
from typing import List
from app.models.schemas import ConversationCreate, ConversationResponse
from app.services.memory import memory_service

router = APIRouter()

@router.get("", response_model=List[ConversationResponse])
async def list_conversations():
    return [] # TODO

@router.get("/{conv_id}")
async def get_conversation(conv_id: str):
    history = await memory_service.get_history(conv_id)
    return {"id": conv_id, "messages": history}

@router.delete("/{conv_id}")
async def delete_conversation(conv_id: str):
    return {"status": "deleted"} # TODO

@router.put("/{conv_id}")
async def update_conversation(conv_id: str, data: dict):
    return {"status": "updated"} # TODO
    """,
    "app/core/__init__.py": "",
    "app/core/llm/__init__.py": "",
    "app/core/llm/router.py": """
import litellm
from typing import AsyncGenerator, List, Dict
from app.config import settings

litellm.drop_params = True

class LLMRouter:
    def __init__(self):
        self.default_model = settings.default_model
        self.vision_model = settings.vision_model

    async def stream_completion(self, messages: List[Dict], model: str = None) -> AsyncGenerator[str, None]:
        target_model = model or self.default_model
        response = await litellm.acompletion(
            model=target_model,
            messages=messages,
            stream=True
        )
        async for chunk in response:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content

    async def completion(self, messages: List[Dict], model: str = None) -> str:
        target_model = model or self.default_model
        response = await litellm.acompletion(
            model=target_model,
            messages=messages
        )
        return response.choices[0].message.content

    async def vision_completion(self, messages: List[Dict], images: List[str]) -> str:
        # Simplification of vision injection
        return await self.completion(messages, model=self.vision_model)

llm_router = LLMRouter()
    """,
    "app/core/llm/prompts.py": """
RAYMUNDO_SYSTEM_PROMPT_AMIGABLE = \"\"\"Eres Raymundo, un asistente de IA creado por Axoloit. Eres amigable, servicial y te expresas con naturalidad. Usas emojis ocasionalmente. Respondes siempre en español. Tienes acceso a herramientas para Google Workspace, YouTube, Spotify y búsqueda web.\"\"\"

RAYMUNDO_SYSTEM_PROMPT_DIRECTO = \"\"\"Eres Raymundo, un asistente de IA creado por Axoloit. Eres directo, conciso y profesional. No usas emojis. Vas al grano sin rodeos. Respondes siempre en español.\"\"\"

RAG_CONTEXT_PROMPT = \"\"\"Utiliza el siguiente contexto para responder a la pregunta. Si no sabes la respuesta basándote en el contexto, dilo.
Contexto:
{context}
\"\"\"

TOOL_USE_PROMPT = \"\"\"Tienes acceso a herramientas (MCP). Utilízalas cuando sea necesario para obtener información en tiempo real o realizar acciones en nombre del usuario.\"\"\"
    """,
    "app/core/rag/__init__.py": "",
    "app/core/rag/pipeline.py": """
from app.core.rag.embeddings import embedding_service
from app.core.rag.chunking import document_chunker
from app.core.rag.retriever import hybrid_retriever
from app.core.rag.ingestion import document_ingestion
from loguru import logger
import uuid

class RAGPipeline:
    async def query(self, question: str, user_id: str, top_k: int = 5):
        query_vector = embedding_service.encode_query(question)
        results = await hybrid_retriever.search(query_vector, user_id, top_k)
        return results

    async def ingest_document(self, file_path: str, user_id: str, metadata: dict) -> str:
        text = await document_ingestion.process_file(file_path, metadata.get("filename", ""))
        chunks = document_chunker.chunk_text(text, metadata)
        
        texts_to_embed = [c.content for c in chunks]
        vectors = embedding_service.encode(texts_to_embed)
        
        doc_id = str(uuid.uuid4())
        await hybrid_retriever.upsert(chunks, vectors, user_id)
        
        return doc_id

    async def delete_document(self, document_id: str) -> bool:
        return await hybrid_retriever.delete(document_id)

rag_pipeline = RAGPipeline()
    """,
    "app/core/rag/embeddings.py": """
from sentence_transformers import SentenceTransformer
from app.config import settings

class EmbeddingService:
    def __init__(self):
        self.model_name = settings.embedding_model
        self._model = None

    @property
    def model(self):
        if self._model is None:
            self._model = SentenceTransformer(self.model_name)
        return self._model

    def encode(self, texts: list[str]) -> list[list[float]]:
        return self.model.encode(texts).tolist()

    def encode_query(self, query: str) -> list[float]:
        return self.model.encode([query])[0].tolist()

embedding_service = EmbeddingService()
    """,
    "app/core/rag/chunking.py": """
from dataclasses import dataclass
from typing import List, Dict, Any

@dataclass
class Chunk:
    content: str
    metadata: Dict[str, Any]
    parent_id: str = None
    chunk_index: int = 0

class DocumentChunker:
    def __init__(self, parent_size: int = 1024, child_size: int = 256):
        self.parent_size = parent_size
        self.child_size = child_size

    def chunk_text(self, text: str, metadata: dict) -> List[Chunk]:
        # Simplistic chunking for demonstration
        words = text.split()
        chunks = []
        for i in range(0, len(words), self.child_size):
            chunk_content = " ".join(words[i:i + self.child_size])
            chunks.append(Chunk(content=chunk_content, metadata=metadata, chunk_index=i))
        return chunks

    def chunk_document(self, file_path: str) -> List[Chunk]:
        return []

document_chunker = DocumentChunker()
    """,
    "app/core/rag/retriever.py": """
from qdrant_client import AsyncQdrantClient
from qdrant_client.http.models import Distance, VectorParams, PointStruct
from app.config import settings
from loguru import logger
import uuid

class HybridRetriever:
    def __init__(self):
        self.client = AsyncQdrantClient(url=settings.qdrant_url)
        self.collection_name = "raygpt_documents"
        self._initialized = False

    async def _init_collection(self):
        if not self._initialized:
            collections = await self.client.get_collections()
            names = [c.name for c in collections.collections]
            if self.collection_name not in names:
                await self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=384, distance=Distance.COSINE)
                )
            self._initialized = True

    async def search(self, query_vector: list[float], user_id: str, top_k: int = 10):
        await self._init_collection()
        # simplified search without user_id filter for now
        results = await self.client.search(
            collection_name=self.collection_name,
            query_vector=query_vector,
            limit=top_k
        )
        # Mocking return structure to match Chunk
        class MockResult:
            def __init__(self, content):
                self.content = content
        return [MockResult(hit.payload.get("text", "")) for hit in results]

    async def upsert(self, chunks: list, vectors: list[list[float]], user_id: str) -> bool:
        await self._init_collection()
        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=vec,
                payload={"text": chunk.content, "user_id": user_id, **chunk.metadata}
            ) for chunk, vec in zip(chunks, vectors)
        ]
        await self.client.upsert(collection_name=self.collection_name, points=points)
        return True

    async def delete(self, document_id: str) -> bool:
        # Placeholder for delete
        return True

hybrid_retriever = HybridRetriever()
    """,
    "app/core/rag/ingestion.py": """
from enum import Enum
import PyPDF2
from docx import Document

class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    UNKNOWN = "unknown"

class DocumentIngestion:
    def get_file_type(self, filename: str) -> FileType:
        ext = filename.split('.')[-1].lower()
        if ext == 'pdf': return FileType.PDF
        elif ext == 'docx': return FileType.DOCX
        elif ext in ['txt', 'md']: return FileType.TXT
        return FileType.UNKNOWN

    async def process_file(self, file_path: str, filename: str) -> str:
        ftype = self.get_file_type(filename)
        text = ""
        
        try:
            if ftype == FileType.PDF:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\\n"
            elif ftype == FileType.DOCX:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\\n"
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
        except Exception as e:
            text = f"Error extracting text: {str(e)}"
            
        return text

document_ingestion = DocumentIngestion()
    """,
    "app/core/mcp/__init__.py": "",
    "app/core/mcp/client_manager.py": """
from loguru import logger

class MCPClientManager:
    def __init__(self):
        self.servers = {}

    async def initialize(self):
        logger.info("Initializing MCP Client Manager...")
        # TODO: Implement full MCP SDK integration in Phase 4

    async def list_all_tools(self) -> dict:
        return {"mock_server": ["search_web", "play_spotify"]}

    async def call_tool(self, server_name: str, tool_name: str, arguments: dict):
        return {"status": "success", "data": f"Mock result for {tool_name}"}

    async def shutdown(self):
        logger.info("Shutting down MCP Client Manager...")

mcp_manager = MCPClientManager()
    """,
    "app/models/__init__.py": "",
    "app/models/schemas.py": """
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from datetime import datetime

class MessageCreate(BaseModel):
    content: str
    role: str
    conversation_id: str
    attachments: Optional[List[str]] = None

class MessageResponse(BaseModel):
    id: str
    content: str
    role: str
    timestamp: datetime
    tool_calls: Optional[List[Dict[str, Any]]] = None

class ConversationCreate(BaseModel):
    title: str
    mode: str

class ConversationResponse(BaseModel):
    id: str
    title: str
    mode: str
    created_at: datetime
    updated_at: datetime
    message_count: int

class DocumentUploadResponse(BaseModel):
    id: str
    filename: str
    file_type: str
    chunk_count: int
    created_at: str

class DocumentResponse(BaseModel):
    id: str
    filename: str
    file_type: str
    chunk_count: int
    created_at: str

class WebSocketMessage(BaseModel):
    type: str
    content: str
    conversation_id: str
    user_id: str
    mode: str

class WebSocketResponse(BaseModel):
    type: str
    content: str
    message_id: Optional[str] = None
    tool_name: Optional[str] = None
    tool_status: Optional[str] = None

class HealthResponse(BaseModel):
    status: str
    version: str
    """,
    "app/db/__init__.py": "",
    "app/db/database.py": """
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.orm import declarative_base, Mapped, mapped_column
from sqlalchemy import String, DateTime, Integer, JSON
from app.config import settings
import datetime

engine = create_async_engine(settings.database_url, echo=False)
AsyncSessionLocal = async_sessionmaker(engine, expire_on_commit=False, class_=AsyncSession)

Base = declarative_base()

class User(Base):
    __tablename__ = "users"
    id: Mapped[str] = mapped_column(String, primary_key=True)
    username: Mapped[str] = mapped_column(String)

class Conversation(Base):
    __tablename__ = "conversations"
    id: Mapped[str] = mapped_column(String, primary_key=True)
    user_id: Mapped[str] = mapped_column(String)
    title: Mapped[str] = mapped_column(String)
    mode: Mapped[str] = mapped_column(String)
    created_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=datetime.datetime.utcnow)
    updated_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)

class Message(Base):
    __tablename__ = "messages"
    id: Mapped[str] = mapped_column(String, primary_key=True)
    conversation_id: Mapped[str] = mapped_column(String)
    content: Mapped[str] = mapped_column(String)
    role: Mapped[str] = mapped_column(String)
    timestamp: Mapped[datetime.datetime] = mapped_column(DateTime, default=datetime.datetime.utcnow)
    tool_calls_json: Mapped[dict] = mapped_column(JSON, nullable=True)

class Document(Base):
    __tablename__ = "documents"
    id: Mapped[str] = mapped_column(String, primary_key=True)
    user_id: Mapped[str] = mapped_column(String)
    filename: Mapped[str] = mapped_column(String)
    file_type: Mapped[str] = mapped_column(String)
    chunk_count: Mapped[int] = mapped_column(Integer, default=0)
    created_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=datetime.datetime.utcnow)
    qdrant_collection: Mapped[str] = mapped_column(String, nullable=True)

async def init_db():
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

async def get_session() -> AsyncSession:
    async with AsyncSessionLocal() as session:
        yield session
    """,
    "app/db/redis_client.py": """
import redis.asyncio as redis
from app.config import settings
from loguru import logger

class RedisClient:
    def __init__(self):
        self.redis = None

    async def init(self):
        self.redis = redis.from_url(settings.redis_url, encoding="utf-8", decode_responses=True)
        logger.info("Redis connected.")

    async def close(self):
        if self.redis:
            await self.redis.close()

    async def cache_set(self, key: str, value: str, expire: int = 3600):
        if self.redis:
            await self.redis.set(key, value, ex=expire)

    async def cache_get(self, key: str) -> str:
        if self.redis:
            return await self.redis.get(key)
        return None

    async def cache_delete(self, key: str):
        if self.redis:
            await self.redis.delete(key)

redis_client = RedisClient()

async def init_redis():
    await redis_client.init()

async def close_redis():
    await redis_client.close()
    """,
    "app/services/__init__.py": "",
    "app/services/memory.py": """
from app.db.database import AsyncSessionLocal, Message, Conversation
from sqlalchemy import select, desc
import uuid
from loguru import logger

class ConversationMemory:
    async def get_history(self, conversation_id: str, limit: int = 50):
        async with AsyncSessionLocal() as session:
            stmt = select(Message).where(Message.conversation_id == conversation_id).order_by(desc(Message.timestamp)).limit(limit)
            result = await session.execute(stmt)
            messages = result.scalars().all()
            return list(reversed(messages))

    async def save_message(self, conversation_id: str, role: str, content: str, tool_calls=None):
        async with AsyncSessionLocal() as session:
            msg = Message(
                id=str(uuid.uuid4()),
                conversation_id=conversation_id,
                content=content,
                role=role,
                tool_calls_json=tool_calls
            )
            session.add(msg)
            await session.commit()
            await session.refresh(msg)
            return msg

    async def create_conversation(self, user_id: str, title: str, mode: str) -> str:
        async with AsyncSessionLocal() as session:
            conv = Conversation(
                id=str(uuid.uuid4()),
                user_id=user_id,
                title=title,
                mode=mode
            )
            session.add(conv)
            await session.commit()
            return conv.id

memory_service = ConversationMemory()
    """
}

def create_files():
    for rel_path, content in FILES.items():
        full_path = os.path.join(BASE_DIR, rel_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        with open(full_path, "w", encoding="utf-8") as f:
            f.write(content.strip() + "\\n")
    print(f"Successfully generated {len(FILES)} files in {BASE_DIR}")

if __name__ == "__main__":
    create_files()
